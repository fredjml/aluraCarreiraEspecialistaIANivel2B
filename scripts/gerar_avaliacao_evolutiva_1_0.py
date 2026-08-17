"""Gera a Avaliação Evolutiva 1.0 do GeoAI Mentor."""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


DESTINO = Path(__file__).parents[1] / "Analise" / "AvaliacaoEvolutiva_1.0_GeoAI_Mentor.docx"
AZUL = "1F4D78"
AZUL_CLARO = "E8EEF5"
PRETO = RGBColor(32, 42, 48)


def fonte(run, tamanho=11, negrito=False, cor=None, italico=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(tamanho)
    run.bold = negrito
    run.italic = italico
    if cor:
        run.font.color.rgb = RGBColor.from_string(cor)


def sombrear(celula, cor):
    propriedades = celula._tc.get_or_add_tcPr()
    shd = propriedades.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        propriedades.append(shd)
    shd.set(qn("w:fill"), cor)


def margens_celula(celula):
    tc_pr = celula._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for lado, valor in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
        item = tc_mar.find(qn(f"w:{lado}"))
        if item is None:
            item = OxmlElement(f"w:{lado}")
            tc_mar.append(item)
        item.set(qn("w:w"), str(valor))
        item.set(qn("w:type"), "dxa")


def tabela(doc, cabecalhos, linhas, larguras):
    tab = doc.add_table(rows=1, cols=len(cabecalhos))
    tab.style = "Table Grid"
    tab.autofit = False
    cabecalho_pr = tab.rows[0]._tr.get_or_add_trPr()
    repetir = OxmlElement("w:tblHeader")
    repetir.set(qn("w:val"), "true")
    cabecalho_pr.append(repetir)
    for i, texto in enumerate(cabecalhos):
        cel = tab.rows[0].cells[i]
        cel.width = Inches(larguras[i])
        cel.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        sombrear(cel, AZUL)
        margens_celula(cel)
        p = cel.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fonte(p.add_run(texto), 10, True, "FFFFFF")
    for indice, linha in enumerate(linhas):
        row = tab.add_row()
        row_pr = row._tr.get_or_add_trPr()
        row_pr.append(OxmlElement("w:cantSplit"))
        cels = row.cells
        for i, texto in enumerate(linha):
            cels[i].width = Inches(larguras[i])
            cels[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            margens_celula(cels[i])
            if indice % 2:
                sombrear(cels[i], "F5F7F9")
            p = cels[i].paragraphs[0]
            fonte(p.add_run(str(texto)), 9.5, False, None)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return tab


def bullet(doc, texto):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    fonte(p.add_run(texto))


def numero(doc, texto):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(5)
    fonte(p.add_run(texto))


def paragrafo(doc, texto, negrito_inicial=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    if negrito_inicial and texto.startswith(negrito_inicial):
        fonte(p.add_run(negrito_inicial), negrito=True)
        fonte(p.add_run(texto[len(negrito_inicial):]))
    else:
        fonte(p.add_run(texto))
    return p


def configurar(doc):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = sec.bottom_margin = Inches(1)
    sec.left_margin = sec.right_margin = Inches(1)
    sec.header_distance = sec.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for nome, tamanho, antes, depois, cor in (
        ("Heading 1", 16, 16, 8, "2E74B5"),
        ("Heading 2", 13, 12, 6, "2E74B5"),
        ("Heading 3", 12, 8, 4, "1F4D78"),
    ):
        estilo = doc.styles[nome]
        estilo.font.name = "Calibri"
        estilo.font.size = Pt(tamanho)
        estilo.font.bold = True
        estilo.font.color.rgb = RGBColor.from_string(cor)
        estilo.paragraph_format.space_before = Pt(antes)
        estilo.paragraph_format.space_after = Pt(depois)
        estilo.paragraph_format.keep_with_next = True


def cabecalho_rodape(doc):
    sec = doc.sections[0]
    p = sec.header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fonte(p.add_run("GeoAI Mentor  |  Avaliação evolutiva 1.0"), 8, False, "687781")
    p = sec.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fonte(p.add_run("Plano incremental de arquitetura e qualidade"), 8, False, "687781")


def gerar():
    doc = Document()
    configurar(doc)
    cabecalho_rodape(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(28)
    p.paragraph_format.space_after = Pt(4)
    fonte(p.add_run("AVALIAÇÃO EVOLUTIVA 1.0"), 23, True, AZUL)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    fonte(p.add_run("GeoAI Mentor - plano gradual de arquitetura, persistência, testes, sessões e RAG"), 13, False, "4E626F")
    for rotulo, valor in (
        ("Data", "17/08/2026"),
        ("Situação analisada", "Interface Streamlit com LangChain e memória temporária"),
        ("Objetivo", "Definir uma evolução segura, testável e incremental"),
        ("Decisão", "Separar camadas antes de persistência; sessões e RAG somente depois dos portões de qualidade"),
    ):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        fonte(p.add_run(f"{rotulo}: "), 10.5, True)
        fonte(p.add_run(valor), 10.5)

    doc.add_heading("1. Conclusão da análise", level=1)
    paragrafo(doc, "A evolução deve ocorrer em blocos com portões de qualidade. A recomendação é não adicionar persistência, sessões avançadas ou RAG sobre a arquitetura ainda acoplada. Primeiro, a interface deve depender apenas de um serviço de aplicação; depois, o armazenamento deve ser introduzido por contrato; por fim, a cobertura deve comprovar a estabilidade antes das fases mais complexas.")
    tabela(doc, ["Ordem", "Evolução", "Portão"], [
        ("1", "Estabilizar a versão atual", "Linha de base reproduzível"),
        ("2", "Separar front-end e back-end", "Interface sem dependência direta de LangChain/OpenAI"),
        ("3", "Introduzir persistência", "Conversas sobrevivem ao reinício"),
        ("4", "Ampliar testes e cobertura", "Cobertura geral igual ou superior a 85%"),
        ("5", "Gerenciar sessões completas", "Ciclo criar, listar, abrir e excluir validado"),
        ("6", "Implementar RAG", "Fontes, segurança e avaliação definidas"),
    ], [0.65, 3.55, 2.3])

    doc.add_heading("2. Diagnóstico atual", level=1)
    paragrafo(doc, "O projeto possui uma interface Streamlit e um núcleo conversacional funcional, mas a separação ainda é parcial. O front-end conhece funções internas do LangChain, o histórico é duplicado entre o componente de memória e o st.session_state, o modelo está fixo no código e ainda não existe uma abstração de armazenamento.")
    for texto in (
        "O dicionário global de memória não é adequado para múltiplos processos.",
        "O encerramento do processo elimina o histórico.",
        "Exceções técnicas podem chegar diretamente ao usuário.",
        "Os testes validam o comportamento básico, mas ainda não medem cobertura.",
        "A CLI e a interface ainda não consomem um serviço de aplicação único.",
    ):
        bullet(doc, texto)

    doc.add_heading("3. Fase 0 - Linha de base", level=1)
    for texto in (
        "Registrar compilação, testes, inicialização do Streamlit e comportamento da memória.",
        "Preservar o .env fora do Git e confirmar que nenhum segredo aparece em saídas.",
        "Resolver separadamente o estado do índice Git antes de qualquer commit, pois existem alterações e arquivos do usuário já preparados.",
        "Estabelecer configuração centralizada e mensagens de erro seguras.",
    ):
        numero(doc, texto)
    paragrafo(doc, "Critério de aceite: testes atuais aprovados, código compilável, segredos protegidos e versão de referência identificável.", "Critério de aceite:")

    doc.add_heading("4. Fase 1 - Separar front-end e back-end", level=1)
    tabela(doc, ["Camada", "Responsabilidade"], [
        ("interfaces", "Entrada, apresentação e estado estritamente visual."),
        ("application", "Coordenação dos casos de uso por meio do MentorService."),
        ("domain", "Modelos e contratos sem dependência de Streamlit, OpenAI ou banco."),
        ("infrastructure", "LangChain, OpenAI e implementações de memória/armazenamento."),
        ("config", "Leitura e validação centralizada das variáveis de ambiente."),
    ], [1.45, 5.05])
    paragrafo(doc, "O front-end deverá chamar apenas um serviço simples, por exemplo enviar_mensagem(session_id, mensagem). A CLI deverá usar o mesmo serviço. LangChain e ChatOpenAI ficarão restritos à infraestrutura, permitindo testes com implementações falsas.")
    paragrafo(doc, "Critério de aceite: o Streamlit não importa LangChain ou ChatOpenAI; o serviço é testável sem interface; configuração e erros possuem fronteiras claras.", "Critério de aceite:")

    doc.add_heading("5. Fase 2 - Persistência", level=1)
    paragrafo(doc, "Recomenda-se SQLite na primeira versão, protegido por uma interface de repositório que permita futura migração para PostgreSQL. O banco será a fonte oficial das mensagens; o st.session_state manterá apenas identificadores e estado temporário da tela.")
    tabela(doc, ["Entidade", "Dados principais"], [
        ("Conversa", "ID, título, criação e atualização."),
        ("Mensagem", "ID, conversa, papel, conteúdo e horário."),
        ("Configuração", "Modelo, temperatura e versão do prompt."),
        ("Evento técnico", "Tipo, situação e descrição sem segredo."),
    ], [1.5, 5.0])
    for texto in (
        "A conversa permanece após reiniciar a aplicação.",
        "Conversas distintas não compartilham conteúdo.",
        "Falhas transacionais não deixam mensagens parciais.",
        "A chave da API nunca é armazenada no banco.",
    ):
        bullet(doc, texto)

    doc.add_heading("6. Fase 3 - Cobertura e qualidade", level=1)
    tabela(doc, ["Métrica", "Meta"], [
        ("Cobertura geral", ">= 85%"),
        ("Serviços, domínio e persistência", ">= 90%"),
        ("Caminhos críticos", "100%"),
        ("Testes que consomem API por padrão", "0"),
    ], [4.7, 1.8])
    for texto in (
        "Testes unitários de configuração, validação, prompt, erros e repositórios.",
        "Testes de integração do serviço com SQLite temporário e pipeline falso.",
        "Testes do Streamlit para envio, resposta, nova conversa e recuperação.",
        "Testes reais da API marcados e opcionais, com consumo mínimo e sem exposição de chave.",
    ):
        bullet(doc, texto)

    heading = doc.add_heading("7. Etapas posteriores - Sessões e RAG", level=1)
    heading.paragraph_format.page_break_before = True
    doc.add_heading("7.1 Sessões completas", level=2)
    paragrafo(doc, "Somente após os portões anteriores: criar, listar, reabrir, renomear e excluir conversas; definir retenção; separar dados por usuário; e validar múltiplas abas e concorrência.")
    doc.add_heading("7.2 RAG", level=2)
    paragrafo(doc, "Antes da implementação, definir fontes aceitas, formatos, atualização, exclusão, confidencialidade, citações e métricas. O primeiro piloto deverá usar um conjunto documental controlado, metadados de fonte e recusa quando não houver evidência.")

    doc.add_heading("8. Portões de avanço", level=1)
    tabela(doc, ["Portão", "Condição obrigatória"], [
        ("Arquitetura", "Front-end sem dependência direta de LangChain/OpenAI."),
        ("Persistência", "Conversas sobrevivem ao reinício e permanecem isoladas."),
        ("Qualidade", "Cobertura mínima de 85% e caminhos críticos aprovados."),
        ("Sessões", "Ciclo criar, listar, abrir e excluir validado."),
        ("RAG", "Fontes, segurança, métricas e conjunto de avaliação definidos."),
    ], [1.4, 5.1])
    paragrafo(doc, "Recomendação executiva: executar agora somente a Fase 0 e a Fase 1. Persistência entrará após a separação estar estável; sessões completas e RAG continuarão fora do escopo até a aprovação dos respectivos portões.", "Recomendação executiva:")

    doc.add_heading("9. Registro de execução desta rodada", level=1)
    paragrafo(doc, "Após o registro do plano, foram executadas somente as fases autorizadas pela recomendação: estabilização da linha de base e separação arquitetural. Persistência, gerenciamento avançado de sessões e RAG não foram iniciados.")
    tabela(doc, ["Verificação", "Resultado"], [
        ("Linha de base anterior", "Compilação aprovada e 6 testes aprovados."),
        ("Separação", "Camadas config, domain, application, infrastructure e interfaces criadas."),
        ("Fronteira do front-end", "Streamlit chama MentorService e não importa LangChain ou ChatOpenAI."),
        ("Compatibilidade", "chatbot_mentor.py e streamlit_app.py permanecem como pontos de entrada."),
        ("Configuração", "Modelo e temperatura centralizados, com validação de ambiente."),
        ("Erros", "Detalhes internos não são mais apresentados ao usuário."),
        ("Suíte final", "16 testes locais aprovados, sem chamada real à API."),
    ], [2.0, 4.5])
    paragrafo(doc, "Situação do Git: nenhum commit foi criado nesta rodada, pois o índice já continha alterações e arquivos preparados pelo usuário. Essa decisão preserva o trabalho existente e evita misturar escopos.", "Situação do Git:")
    paragrafo(doc, "Próximo portão: revisar e consolidar o estado do Git; depois planejar a persistência SQLite por repositório, sem ainda iniciar sessões avançadas ou RAG.", "Próximo portão:")

    doc.add_heading("10. Registro da Fase 2 - Persistência SQLite", level=1)
    paragrafo(doc, "A persistência foi implementada após a aprovação da separação arquitetural. O banco SQLite passou a ser a fonte oficial do histórico, enquanto o st.session_state permaneceu limitado ao estado visual da interface.")
    tabela(doc, ["Componente", "Implementação e evidência"], [
        ("Contrato", "ConversationRepository definido no domínio, sem dependência de SQLite."),
        ("Modelo", "Message registra papel, conteúdo e instante de criação."),
        ("SQLite", "Tabelas conversations e messages, chave estrangeira, exclusão em cascata e índice de ordenação."),
        ("Atomicidade", "Pergunta e resposta são gravadas na mesma transação; falha simulada reverteu as duas."),
        ("Recuperação", "Uma nova instância do repositório recuperou a interação gravada em arquivo."),
    ], [1.65, 4.85])
    heading = doc.add_heading("10. Registro da Fase 2 - continuação", level=2)
    tabela(doc, ["Componente", "Implementação e evidência"], [
        ("Isolamento", "Conversas com identificadores diferentes mantiveram mensagens independentes."),
        ("LangChain", "O gateway carrega mensagens do repositório e não utiliza mais RunnableWithMessageHistory."),
        ("Configuração", "GEOAI_DATABASE_PATH permite escolher o banco; o padrão é data/geoai_mentor.db."),
        ("Proteção Git", "Arquivos .db, .db-shm e .db-wal do diretório data estão ignorados."),
        ("Suíte", "20 testes aprovados, sem chamada real à API e sem aviso de depreciação do mecanismo anterior."),
    ], [1.65, 4.85])
    doc.add_heading("10.1 Limite deliberado da fase", level=2)
    paragrafo(doc, "O conteúdo já sobrevive ao encerramento do processo quando o mesmo identificador é reutilizado. A interface ainda não lista nem reabre conversas anteriores, pois esse comportamento pertence à futura fase de gerenciamento de sessões e depende primeiro da aprovação da cobertura.")
    doc.add_heading("10.2 Próximo portão", level=2)
    paragrafo(doc, "Instalar e configurar pytest-cov, medir a cobertura real, elevar serviços, domínio e persistência aos limites definidos e tornar a verificação reproduzível. Sessões avançadas e RAG permanecem bloqueados.")

    DESTINO.parent.mkdir(exist_ok=True)
    doc.save(DESTINO)


if __name__ == "__main__":
    gerar()
