"""Gera DOCX profissionais a partir dos relatórios Markdown versionados."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "Docs"
REPORTS = (
    ("relatorio_levantamento_bytebank.md", "Levantamento", "Requisitos, lacunas, decisões e controles"),
    ("relatorio_implementacao_bytebank.md", "Implementação", "Evidências técnicas e prontidão para avaliação"),
    ("relatorio_executivo_implementacao_bytebank.md", "Executivo", "Visão de resultado, risco e maturidade técnica"),
)
NAVY = "0F3557"
BLUE = RGBColor(15, 90, 138)
LIGHT_BLUE = "EAF4FB"
LIGHT_GRAY = "F3F6F8"


def shade(cell, color: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    fill = OxmlElement("w:shd")
    fill.set(qn("w:fill"), color)
    properties.append(fill)


def set_cell_margins(cell, top=90, start=100, bottom=90, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        margins.append(node)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("PÁGINA ")
    run.font.size = Pt(8)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def configure_document(document: Document, subtitle: str) -> None:
    section = document.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(38, 58, 73)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.12

    for name, size in (("Title", 31), ("Heading 1", 19), ("Heading 2", 14), ("Heading 3", 11)):
        style = styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.color.rgb = BLUE
        style.font.bold = True
        style.paragraph_format.space_before = Pt(12 if name != "Title" else 0)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True

    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.text = f"BYTEBANK  /  AI ECOSYSTEM     {subtitle.upper()}"
    p.style = styles["Caption"]
    p.runs[0].font.color.rgb = BLUE
    p.runs[0].font.bold = True
    p.runs[0].font.size = Pt(8)
    add_page_number(section.footer.paragraphs[0])


def add_cover(document: Document, report_name: str, tagline: str) -> None:
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(28)
    band = document.add_table(rows=1, cols=1)
    band.autofit = False
    band.columns[0].width = Inches(6.7)
    cell = band.cell(0, 0)
    shade(cell, NAVY)
    set_cell_margins(cell, top=520, start=420, bottom=520, end=420)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    eyebrow = p.add_run("BYTEBANK  ·  ESPECIALISTA EM IA NÍVEL 2\n")
    eyebrow.font.name = "Aptos"
    eyebrow.font.size = Pt(10)
    eyebrow.font.bold = True
    eyebrow.font.color.rgb = RGBColor(163, 214, 242)
    title = p.add_run(f"Relatório de {report_name}\n")
    title.font.name = "Aptos Display"
    title.font.size = Pt(31)
    title.font.bold = True
    title.font.color.rgb = RGBColor(255, 255, 255)
    sub = p.add_run(tagline)
    sub.font.name = "Aptos"
    sub.font.size = Pt(13)
    sub.font.color.rgb = RGBColor(231, 244, 251)

    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.add_run("DESAFIO DE PORTFÓLIO\n").bold = True
    p.add_run("Governança · RAG · ChromaDB · Multiagentes · A2A · MCP · HITL\n")
    date_run = p.add_run("22 de agosto de 2026")
    date_run.font.color.rgb = BLUE
    document.add_paragraph(
        "Documento de projeto fictício. Não contém dados reais de clientes nem credenciais."
    ).style = document.styles["Caption"]
    document.add_page_break()


def clean_inline(text: str) -> str:
    return text.replace("**", "").replace("`", "")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def add_markdown_table(document: Document, lines: list[str]) -> None:
    raw_rows = [[clean_inline(cell.strip()) for cell in line.strip().strip("|").split("|")] for line in lines]
    rows = [raw_rows[0]] + raw_rows[2:]
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.autofit = True
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = value
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(1)
                for run in paragraph.runs:
                    run.font.size = Pt(8.5)
            if r_idx == 0:
                shade(cell, NAVY)
                for run in cell.paragraphs[0].runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.bold = True
            elif r_idx % 2 == 0:
                shade(cell, LIGHT_BLUE)
    set_repeat_table_header(table.rows[0])
    document.add_paragraph().paragraph_format.space_after = Pt(1)


def add_code_line(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.left_indent = Cm(0.35)
    shade_proxy = OxmlElement("w:shd")
    shade_proxy.set(qn("w:fill"), LIGHT_GRAY)
    p._p.get_or_add_pPr().append(shade_proxy)
    run = p.add_run(text or " ")
    run.font.name = "Cascadia Mono"
    run.font.size = Pt(8.5)


def markdown_to_docx(markdown_path: Path, output_path: Path, report_name: str, tagline: str) -> None:
    text = markdown_path.read_text(encoding="utf-8")
    document = Document()
    configure_document(document, report_name)
    add_cover(document, report_name, tagline)
    document.core_properties.title = f"Relatório de {report_name} · Bytebank Nível 2"
    document.core_properties.subject = "Entregáveis do desafio Especialista em IA Nível 2"
    document.core_properties.author = "Projeto Bytebank AI Ecosystem"

    lines = text.splitlines()
    index = 1 if lines and lines[0].startswith("# ") else 0
    in_code = False
    while index < len(lines):
        line = lines[index]
        if line.startswith("```"):
            in_code = not in_code
            index += 1
            continue
        if in_code:
            add_code_line(document, line)
            index += 1
            continue
        if line.startswith("|") and index + 1 < len(lines) and re.match(r"^\|[-:| ]+\|$", lines[index + 1]):
            table_lines = [line, lines[index + 1]]
            index += 2
            while index < len(lines) and lines[index].startswith("|"):
                table_lines.append(lines[index])
                index += 1
            add_markdown_table(document, table_lines)
            continue
        if not line.strip():
            index += 1
            continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading:
            level = len(heading.group(1))
            document.add_heading(clean_inline(heading.group(2)), level=level)
        elif re.match(r"^\d+\.\s+", line):
            document.add_paragraph(clean_inline(re.sub(r"^\d+\.\s+", "", line)), style="List Number")
        elif line.startswith("- "):
            document.add_paragraph(clean_inline(line[2:]), style="List Bullet")
        elif line.startswith("> "):
            p = document.add_paragraph(clean_inline(line[2:]))
            p.style = document.styles["Quote"]
        else:
            p = document.add_paragraph(clean_inline(line))
            if line.startswith("**"):
                for run in p.runs:
                    run.bold = True
        index += 1

    document.save(output_path)


def main() -> None:
    for source, name, tagline in REPORTS:
        output = DOCS / source.replace(".md", ".docx")
        markdown_to_docx(DOCS / source, output, name, tagline)
        print(f"Gerado: {output}")


if __name__ == "__main__":
    main()
