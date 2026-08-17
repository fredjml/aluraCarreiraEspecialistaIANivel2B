"""Registra a evolução Streamlit nos documentos existentes do projeto."""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
from docx.shared import Inches


RAIZ = Path(__file__).parents[1]
PASSO_A_PASSO = RAIZ / "Analise" / "RegistroPassoAPasso_Implementacao_GeoAI_Mentor.docx"
EXECUTIVO = RAIZ / "Analise" / "RelatorioExecutivoConsolidado_GeoAI_Mentor.docx"


def adicionar_tabela(documento, cabecalhos, linhas, larguras):
    tabela = documento.add_table(rows=1, cols=len(cabecalhos))
    tabela.style = "Table Grid"
    tabela.autofit = False
    for indice, texto in enumerate(cabecalhos):
        celula = tabela.rows[0].cells[indice]
        celula.width = Inches(larguras[indice])
        celula.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragrafo = celula.paragraphs[0]
        paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragrafo.add_run(texto)
        run.bold = True
    for linha in linhas:
        celulas = tabela.add_row().cells
        for indice, texto in enumerate(linha):
            celulas[indice].width = Inches(larguras[indice])
            celulas[indice].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            celulas[indice].text = texto
    return tabela


def atualizar_passo_a_passo() -> None:
    doc = Document(PASSO_A_PASSO)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    doc.add_heading("12. Etapa 8 - Interface Streamlit e testes automatizados", level=1)
    doc.add_paragraph(
        "Data da implementação: 17/08/2026. Objetivo: transformar o protótipo de "
        "terminal em uma experiência de chat amigável, preservando o modelo OpenAI, "
        "a persona especializada e a memória por sessão."
    )

    doc.add_heading("12.1 Itens implementados", level=2)
    adicionar_tabela(
        doc,
        ["Item", "Implementação", "Resultado"],
        [
            ("1. Interface", "streamlit_app.py com mensagens, entrada de chat, barra lateral e botão Nova conversa.", "Concluído"),
            ("2. Histórico visual", "st.session_state armazena o identificador exclusivo e as mensagens exibidas no navegador.", "Concluído"),
            ("3. Entrada dinâmica", "Cada texto digitado é enviado pela função perguntar() ao fluxo LangChain com o session_id correto.", "Concluído"),
            ("6. Portfólio", "README, testes, GIF demonstrativo e documentos foram atualizados.", "Concluído"),
        ],
        [1.2, 4.4, 1.0],
    )

    doc.add_heading("12.2 Decisões de arquitetura", level=2)
    for texto in [
        "A memória do LangChain continua responsável pelo contexto enviado ao modelo.",
        "O st.session_state conserva apenas o estado da experiência visual em cada navegador.",
        "Um UUID separa as conversas; Nova conversa remove o histórico anterior e cria outro identificador.",
        "A cadeia é armazenada com st.cache_resource para evitar reconstrução desnecessária a cada atualização da tela.",
        "A credencial permanece somente no .env e não é exibida na interface, nos testes ou no GIF.",
    ]:
        doc.add_paragraph(texto, style="List Bullet")

    doc.add_page_break()
    doc.add_heading("12.3 Evidências técnicas", level=2)
    adicionar_tabela(
        doc,
        ["Verificação", "Comando / mecanismo", "Resultado observado"],
        [
            ("Compilação", "python -m py_compile chatbot_mentor.py streamlit_app.py", "Aprovada, sem erros"),
            ("Testes", "python -m pytest -q", "6 testes aprovados"),
            ("Memória", "RunnableLambda sem API", "Contexto preservado e sessões isoladas"),
            ("Interface", "streamlit.testing.v1.AppTest", "Título, mensagem, campo e reinício confirmados"),
            ("GIF", "scripts/gerar_gif_demo.py", "4 quadros, 1100 x 680 pixels"),
        ],
        [1.25, 2.8, 2.55],
    )

    doc.add_heading("12.4 Como reproduzir", level=2)
    for texto in [
        "Ative o ambiente virtual com .venv\\Scripts\\Activate.ps1.",
        "Instale a aplicação com pip install -r requirements.txt.",
        "Instale os recursos de teste com pip install -r requirements-dev.txt.",
        "Execute python -m pytest -q e confirme 6 testes aprovados.",
        "Execute streamlit run streamlit_app.py e abra o endereço local informado.",
        "Faça uma pergunta e depois outra que dependa da primeira; use Nova conversa para validar o isolamento.",
    ]:
        doc.add_paragraph(texto, style="List Number")

    doc.add_heading("12.5 Resultado da etapa", level=2)
    doc.add_paragraph(
        "O GeoAI Mentor agora possui uma interface web local adequada a demonstrações e "
        "usuários não técnicos. A entrada deixou de ser uma lista fixa: cada mensagem é "
        "digitada pelo usuário, exibida na tela e vinculada à memória da sessão. A "
        "implementação foi aprovada por testes automatizados sem consumir créditos da API."
    )
    doc.save(PASSO_A_PASSO)


def atualizar_executivo() -> None:
    doc = Document(EXECUTIVO)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    doc.add_heading("Atualização executiva - Interface web", level=1)
    doc.add_paragraph(
        "Em 17/08/2026, o protótipo evoluiu do terminal para uma interface web local em "
        "Streamlit. A evolução preserva a integração com OpenAI e LangChain e acrescenta "
        "uma experiência de conversa mais acessível para pessoas não técnicas."
    )
    adicionar_tabela(
        doc,
        ["Indicador", "Resultado", "Leitura executiva"],
        [
            ("Interface web local", "Entregue", "O usuário conversa pelo navegador, sem editar o código."),
            ("Memória visual", "Entregue", "As mensagens permanecem na tela durante a sessão."),
            ("Isolamento", "Entregue", "Nova conversa inicia um histórico independente."),
            ("Testes automatizados", "6 de 6 aprovados", "Núcleo e interface foram verificados sem custo de API."),
            ("Demonstração", "GIF atualizado", "O portfólio apresenta visualmente o fluxo com memória."),
        ],
        [1.55, 1.3, 3.75],
    )
    doc.add_heading("Impacto para o usuário", level=2)
    doc.add_paragraph(
        "O GeoAI Mentor passa a oferecer um campo de mensagem, histórico visível e um botão "
        "para reiniciar a conversa. A interface reduz a barreira de entrada e torna a demonstração "
        "do portfólio mais clara, mantendo a mesma especialização em geociências e dados."
    )
    doc.add_heading("Limites e próximo passo", level=2)
    doc.add_paragraph(
        "A aplicação ainda é executada localmente e a memória continua temporária. Para uso "
        "mais amplo, recomenda-se hospedar o Streamlit em ambiente controlado, adicionar "
        "autenticação, persistência e limites de consumo da API."
    )
    doc.add_heading("Reprodução resumida", level=2)
    for texto in [
        "Instale as dependências com pip install -r requirements.txt.",
        "Execute os testes com python -m pytest -q.",
        "Inicie a interface com streamlit run streamlit_app.py.",
    ]:
        doc.add_paragraph(texto, style="List Bullet")
    doc.save(EXECUTIVO)


if __name__ == "__main__":
    atualizar_passo_a_passo()
    atualizar_executivo()
