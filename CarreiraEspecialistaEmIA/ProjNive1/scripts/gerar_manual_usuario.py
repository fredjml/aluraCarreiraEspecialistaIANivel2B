from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parents[1] / "Analise" / "ManualDoUsuario_GeoAI_Mentor.docx"
BLUE = "2E74B5"
DARK = "1F4D78"
INK = "0B2545"
LIGHT = "E8EEF5"
PALE = "F4F6F9"
GOLD = "7A5A00"
RED = "9B1C1C"
GRAY = "5B6573"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    set_font(run, 9, GRAY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


ACTIVE_NUM_ID = None


def start_numbered_list(doc):
    global ACTIVE_NUM_ID
    numbering = doc.part.numbering_part.element
    style_num_id = doc.styles["List Number"]._element.pPr.numPr.numId.val
    source_num = next(n for n in numbering.findall(qn("w:num")) if int(n.get(qn("w:numId"))) == style_num_id)
    abstract_id = source_num.find(qn("w:abstractNumId")).get(qn("w:val"))
    new_id = max([int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))] + [0]) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_id)
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    ACTIVE_NUM_ID = new_id


def add_step(doc, text):
    p = doc.add_paragraph(style="List Number")
    num_pr = p._p.get_or_add_pPr().get_or_add_numPr()
    num_pr.get_or_add_ilvl().val = 0
    num_pr.get_or_add_numId().val = ACTIVE_NUM_ID
    p.add_run(text)
    return p


def add_code(doc, lines):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    set_table_geometry(t, [9360])
    cell = t.cell(0, 0)
    shade(cell, "F2F4F7")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for i, line in enumerate(lines):
        if i:
            p.add_run("\n")
        r = p.add_run(line)
        r.font.name = "Consolas"
        r._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Consolas")
        r._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Consolas")
        r.font.size = Pt(9)
    return t


def add_callout(doc, title, text, color=INK, fill=PALE):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    set_table_geometry(t, [9360])
    cell = t.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title + " ")
    set_font(r, 10.5, color, True)
    r = p.add_run(text)
    set_font(r, 10.5, INK)
    return t


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for i, value in enumerate(headers):
        shade(table.rows[0].cells[i], LIGHT)
        p = table.rows[0].cells[i].paragraphs[0]
        r = p.add_run(value)
        set_font(r, 9.5, INK, True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            r = p.add_run(str(value))
            set_font(r, 9.5, INK)
    set_table_geometry(table, widths)
    return table


doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin = sec.bottom_margin = Inches(1)
sec.left_margin = sec.right_margin = Inches(1)
sec.header_distance = sec.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25
for name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 12, DARK, 10, 5),
):
    st = styles[name]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True
for list_name in ("List Bullet", "List Number"):
    st = styles[list_name]
    st.font.name = "Calibri"
    st.font.size = Pt(11)
    st.paragraph_format.left_indent = Inches(0.375)
    st.paragraph_format.first_line_indent = Inches(-0.188)
    st.paragraph_format.space_after = Pt(4)
    st.paragraph_format.line_spacing = 1.25

header = sec.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = header.add_run("GeoAI Mentor  |  Manual do usuário")
set_font(r, 9, GRAY, True)
footer = sec.footer.paragraphs[0]
add_page_number(footer)

# Editorial cover
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(80)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("GUIA DA POC")
set_font(r, 11, GOLD, True)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(8)
r = p.add_run("GeoAI Mentor")
set_font(r, 30, INK, True)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(28)
r = p.add_run("Manual do usuário, testes e operação local")
set_font(r, 15, DARK)
add_callout(doc, "Escopo:", "prova de conceito para dois participantes. O ambiente local está funcional; Entra ID, App Service e PostgreSQL descrevem a próxima implantação e ainda não estão provisionados.", GOLD, "FFF8E8")
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(42)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Versão 1.0  |  17 de agosto de 2026")
set_font(r, 11, GRAY, True)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Início previsto do piloto: 24 de agosto de 2026")
set_font(r, 10, GRAY)
doc.add_page_break()

add_heading(doc, "1. Antes de começar")
doc.add_paragraph("Este manual orienta participantes e responsáveis pela POC na instalação, execução, validação e operação básica do GeoAI Mentor. Ele consolida o registro da implementação evolutiva e o roteiro de testes e custos do ambiente.")
add_callout(doc, "Importante:", "o GeoAI Mentor produz orientação assistida por IA. Não use as respostas como decisão técnica, profissional ou institucional sem revisão humana.", RED, "FDECEC")
add_heading(doc, "O que já funciona", 2)
for item in [
    "Interface web Streamlit e versão de terminal.",
    "Conversas persistidas em SQLite: criar, listar, reabrir, renomear e excluir.",
    "Memória por conversa, com isolamento entre identificadores.",
    "RAG lexical local com dez fontes Markdown curadas em Analise/docsgeo.",
    "Retenção configurável, backup local, limites da API e logs redigidos.",
    "49 testes automatizados aprovados e cobertura total de 88,50%.",
]: add_bullet(doc, item)
add_heading(doc, "O que ainda não está disponível", 2)
for item in [
    "Login pelo Microsoft Entra ID e separação persistente por usuário autenticado.",
    "Publicação no Azure App Service e banco Azure Database for PostgreSQL.",
    "Bloqueio financeiro integral somando Azure e OpenAI.",
    "Piloto real com os dois geocientistas e avaliação final das fontes institucionais.",
]: add_bullet(doc, item)

doc.add_page_break()
add_heading(doc, "2. Preparar o ambiente")
add_heading(doc, "Pré-requisitos", 2)
for item in ["Windows com PowerShell.", "Python 3.12 recomendado.", "Chave válida da API OpenAI para testes com respostas reais.", "Acesso à pasta E:\\ProjAlura."]:
    add_bullet(doc, item)
add_heading(doc, "Configuração inicial", 2)
start_numbered_list(doc)
for step in [
    "Abra o PowerShell na pasta E:\\ProjAlura.",
    "Ative o ambiente virtual.",
    "Instale as dependências de desenvolvimento.",
    "Copie .env.example para .env e preencha somente a sua chave local.",
]: add_step(doc, step)
add_code(doc, [
    "cd E:\\ProjAlura",
    ".venv\\Scripts\\Activate.ps1",
    "pip install -r requirements-dev.txt",
    "Copy-Item .env.example .env",
])
add_heading(doc, "Variáveis principais", 2)
add_table(doc, ["Variável", "Finalidade / padrão"], [
    ("OPENAI_API_KEY", "Obrigatória para respostas reais; nunca versionar."),
    ("OPENAI_MODEL", "Modelo de texto; padrão atual: gpt-5.6-sol."),
    ("GEOAI_DATABASE_PATH", "Banco local; padrão: data/geoai_mentor.db."),
    ("GEOAI_KNOWLEDGE_PATH", "Base RAG; padrão: Analise/docsgeo."),
    ("OPENAI_REQUEST_TIMEOUT", "Tempo limite da chamada; padrão: 30 segundos."),
    ("OPENAI_MAX_OUTPUT_TOKENS", "Máximo de saída; padrão: 1.200 tokens."),
    ("GEOAI_RETENTION_DAYS", "Retenção local; padrão aprovado: 90 dias."),
], [2700, 6660])
add_callout(doc, "Segurança:", "nunca cole a chave em código, captura de tela, relatório, mensagem ou Git. Se houver exposição, revogue-a e crie outra.", RED, "FDECEC")

add_heading(doc, "3. Testar antes de usar")
add_heading(doc, "Teste automatizado", 2)
doc.add_paragraph("A suíte padrão usa substitutos controlados e não deve realizar chamadas reais à OpenAI. Portanto, executar apenas o pytest não gera custo de API.")
add_code(doc, ["python -m pytest -q"])
add_table(doc, ["Critério", "Resultado esperado"], [
    ("Testes", "49 aprovados."),
    ("Cobertura", "Pelo menos 85%; última medição: 88,50%."),
    ("API real", "Nenhuma chamada nos testes padrão."),
    ("Concorrência", "12 gravações em 4 workers, isoladas e completas."),
], [2400, 6960])
add_heading(doc, "Abrir a interface", 2)
add_code(doc, ["streamlit run streamlit_app.py"])
doc.add_paragraph("O navegador deverá abrir automaticamente. Se não abrir, use o endereço local exibido no PowerShell, normalmente http://localhost:8501.")

add_heading(doc, "4. Usar o GeoAI Mentor")
add_heading(doc, "Conversa básica", 2)
start_numbered_list(doc)
for step in [
    "Digite uma pergunta no campo de chat e envie.",
    "Aguarde a resposta; durante chamadas reais pode haver consumo da API.",
    "Faça uma pergunta complementar para verificar se o contexto foi preservado.",
    "Confira as fontes quando a resposta utilizar a base RAG.",
]: add_step(doc, step)
add_callout(doc, "Exemplo de memória:", "pergunte qual linguagem aprender primeiro; depois pergunte 'Que projeto posso criar com essa linguagem?'. A segunda resposta deve compreender a referência anterior.")
p = add_heading(doc, "Gerenciar conversas", 2)
p.paragraph_format.page_break_before = True
add_table(doc, ["Ação", "Como usar"], [
    ("Nova conversa", "Use o comando correspondente na barra lateral."),
    ("Reabrir", "Selecione uma conversa existente; o histórico volta à tela."),
    ("Renomear", "Altere o título para facilitar a localização."),
    ("Excluir", "Confirme a exclusão; as mensagens relacionadas são removidas em cascata."),
], [2000, 7360])
add_heading(doc, "Validar o RAG", 2)
doc.add_paragraph("A recuperação pesquisa somente arquivos Markdown da pasta autorizada. Use perguntas ligadas aos temas dos dez documentos e confirme se a resposta identifica a fonte. Para um assunto sem evidência na base, o comportamento esperado é informar a ausência de suporte, sem inventar uma fonte.")
add_callout(doc, "Limite conhecido:", "a busca é lexical, local e adequada à POC. Os dez documentos ainda precisam de aprovação institucional final antes de uso além do piloto.", GOLD, "FFF8E8")

add_heading(doc, "5. Roteiro de validação da POC")
add_table(doc, ["Etapa", "Ação", "Aceite"], [
    ("1", "Executar pytest.", "49 testes aprovados; cobertura >= 85%."),
    ("2", "Abrir o Streamlit.", "Tela carrega sem revelar segredo."),
    ("3", "Criar duas conversas.", "Históricos não se misturam."),
    ("4", "Reabrir e renomear.", "Conteúdo e novo título permanecem."),
    ("5", "Fazer pergunta coberta pelo RAG.", "Resposta traz fonte pertinente."),
    ("6", "Perguntar fora da base.", "Sistema declara falta de evidência."),
    ("7", "Excluir uma conversa.", "Ela e suas mensagens deixam de aparecer."),
    ("8", "Criar e verificar backup.", "Arquivo restaurável é criado em backups/."),
], [800, 3880, 4680])
add_heading(doc, "Feedback dos dois participantes", 2)
for item in [
    "Utilidade percebida: meta média >= 4/5.",
    "Clareza: meta média >= 4/5.",
    "Fontes corretas e presentes quando necessárias: meta >= 90%.",
    "Recusa adequada fora da base: meta >= 90%.",
    "Conclusão das tarefas: meta >= 80%.",
    "Zero exposição de segredo e zero acesso cruzado.",
]: add_bullet(doc, item)

add_heading(doc, "6. Operação e manutenção local")
add_heading(doc, "Inspecionar o banco", 2)
add_code(doc, ["python scripts\\operacoes_geoai.py status"])
add_heading(doc, "Criar backup", 2)
add_code(doc, ["python scripts\\operacoes_geoai.py backup"])
doc.add_paragraph("O backup consistente é gravado por padrão em backups/ e fica fora do Git. O período aprovado para os backups da futura POC Azure é de 30 dias.")
add_heading(doc, "Aplicar retenção", 2)
add_code(doc, ["python scripts\\operacoes_geoai.py retencao --dias 90"])
add_callout(doc, "Atenção:", "a retenção exclui conversas cuja última atualização seja anterior ao limite. Crie e verifique um backup antes da execução manual.", RED, "FDECEC")
add_heading(doc, "Diagnóstico rápido", 2)
add_table(doc, ["Sintoma", "Verificação"], [
    ("Chave ausente", "Confirme que .env existe e contém OPENAI_API_KEY, sem exibir o valor."),
    ("Modelo não responde", "Verifique internet, saldo, modelo permitido e timeout."),
    ("RAG não encontra fonte", "Confirme GEOAI_KNOWLEDGE_PATH e arquivos .md em Analise/docsgeo."),
    ("Histórico não aparece", "Confirme GEOAI_DATABASE_PATH e permissão de escrita na pasta data."),
    ("Streamlit não abre", "Leia a URL no terminal e confirme que a porta não está ocupada."),
], [2600, 6760])

add_heading(doc, "7. Custos e orçamento da POC")
doc.add_paragraph("O orçamento total aprovado é R$ 50 para Azure e OpenAI em conjunto. Os valores abaixo são estimativas de planejamento e variam com câmbio, impostos, tokens, tempo de execução e preços vigentes.")
add_table(doc, ["Componente / cenário", "Estimativa", "Observação"], [
    ("Testes automatizados", "R$ 0 de API", "Usam substitutos; não chamam a OpenAI."),
    ("App Service Linux F1", "R$ 0", "Sem SLA e com limites; apropriado apenas à POC."),
    ("PostgreSQL B1ms por 14 dias", "~R$ 29,50 + armazenamento", "Parar quando ocioso; armazenamento continua cobrado."),
    ("OpenAI Sol, 20 interações típicas", "~R$ 4", "Exemplo de 3.000 tokens de entrada e 800 de saída."),
    ("Margem de segurança", "R$ 10", "Reserva de 20% do orçamento."),
    ("Total de planejamento", "~R$ 44-50", "Depende sobretudo do banco e do uso real."),
], [3300, 1900, 4160])
add_heading(doc, "Economizar durante o piloto", 2)
for item in [
    "Preferir gpt-5.6-luna quando a qualidade for suficiente; no mesmo exemplo, a chamada seria cerca de R$ 0,04 em vez de R$ 0,20.",
    "Parar o PostgreSQL quando o piloto estiver inativo e remover recursos ao encerrar.",
    "Manter respostas curtas, limitar histórico e acompanhar tokens por participante.",
    "Usar alertas em 50%, 75%, 80%, 90% e 100%; bloquear novas chamadas OpenAI no teto interno de R$ 40.",
]: add_bullet(doc, item)
add_callout(doc, "Azure for Students:", "o crédito pode evitar desembolso direto enquanto houver saldo, mas o consumo precisa continuar registrado no limite da POC.", GOLD, "FFF8E8")

p = add_heading(doc, "8. Evolução por portões")
p.paragraph_format.page_break_before = True
add_table(doc, ["Portão", "Situação", "Entrega principal"], [
    ("0", "Aprovado", "Linha de base e credencial fora do Git."),
    ("1", "Aprovado", "Separação em interface, aplicação, domínio, infraestrutura e configuração."),
    ("2", "Aprovado", "SQLite transacional e isolamento por conversa."),
    ("3", "Aprovado", "Cobertura mínima automatizada e testes sem API real."),
    ("4", "Aprovado", "Gerenciamento completo das conversas."),
    ("5", "Aprovado para piloto", "RAG local controlado, fontes e recusa sem evidência."),
    ("6", "Parcialmente aprovado", "Prontidão técnica local concluída; validações externas pendentes."),
], [1000, 2400, 5960])
add_heading(doc, "Próxima implantação planejada", 2)
add_table(doc, ["Decisão", "Definição da POC"], [
    ("Identidade", "Microsoft Entra ID; tenant único."),
    ("Hospedagem", "Azure App Service, região eastus."),
    ("Persistência", "Azure Database for PostgreSQL, iniciado vazio."),
    ("Participantes", "Dois; contas autorizadas ainda precisam ser definidas."),
    ("Retenção", "Conversas por 90 dias; backups por 30 dias."),
    ("Responsável por alertas", "fredjml.br@gmail.com."),
    ("Início previsto", "24/08/2026."),
], [2700, 6660])
add_callout(doc, "Estado real:", "esta tabela registra decisões, não comprova implantação. Autenticação, PostgreSQL, hospedagem e controle unificado de custos ainda exigem implementação e provisionamento.", GOLD, "FFF8E8")

add_heading(doc, "9. Uso responsável e privacidade")
for item in [
    "Não inserir dados pessoais, confidenciais, estratégicos ou sujeitos a sigilo.",
    "Revisar toda recomendação antes de usá-la em trabalho técnico ou decisão profissional.",
    "Registrar feedback de forma anonimizada, conforme aprovado para a POC.",
    "Excluir conversas de teste que não precisem ser preservadas.",
    "Comunicar imediatamente suspeita de vazamento, acesso indevido ou gasto anormal ao responsável do piloto.",
]: add_bullet(doc, item)
add_heading(doc, "Pendências antes do piloto hospedado", 2)
for item in [
    "Definir as duas contas participantes e se o tenant inteiro ou apenas uma lista terá acesso.",
    "Definir administrador, suporte, privacidade, incidentes e solicitações de exclusão.",
    "Definir acesso administrativo ao PostgreSQL e destino autorizado dos backups.",
    "Aprovar formalmente os dez documentos RAG e a regra de citação.",
    "Definir duração, número de conversas e comportamento exato ao atingir o limite financeiro.",
]: add_bullet(doc, item)

add_heading(doc, "10. Referências")
doc.add_paragraph("Fontes internas consolidadas: RegistroImplementacaoEvolutiva.md, DecisoesDeImplementacao_PortaoIntegral.md, README.md e código operacional do projeto.")
for title, url in [
    ("OpenAI - gpt-5.6-sol", "https://developers.openai.com/api/docs/models/gpt-5.6-sol"),
    ("OpenAI - comparação de modelos", "https://developers.openai.com/api/docs/models/text"),
    ("Azure App Service Linux - preços", "https://azure.microsoft.com/en-us/pricing/details/app-service/linux/"),
    ("Azure - limites dos serviços", "https://learn.microsoft.com/en-us/azure/azure-resource-manager/management/azure-subscription-service-limits"),
    ("Azure Database for PostgreSQL - preços", "https://azure.microsoft.com/pt-br/pricing/details/postgresql/flexible-server/"),
    ("Azure Database for PostgreSQL - visão geral", "https://learn.microsoft.com/en-us/azure/postgresql/overview"),
    ("Azure for Students - acompanhamento de custos", "https://learn.microsoft.com/en-us/azure/education-hub/navigate-costs"),
]:
    p = doc.add_paragraph()
    r = p.add_run(f"{title}: {url}")
    set_font(r, 9.5, DARK)

doc.core_properties.title = "Manual do usuário - GeoAI Mentor"
doc.core_properties.subject = "Uso, testes e operação da prova de conceito"
doc.core_properties.author = "Projeto GeoAI Mentor"
doc.core_properties.keywords = "GeoAI Mentor, POC, manual, Streamlit, RAG, OpenAI, Azure"
OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
