"""Aplica correções locais de paginação aos documentos já atualizados."""

from pathlib import Path

from docx import Document


RAIZ = Path(__file__).parents[1] / "Analise"


passo = RAIZ / "RegistroPassoAPasso_Implementacao_GeoAI_Mentor.docx"
doc = Document(passo)
for paragrafo in doc.paragraphs:
    if paragrafo.text.strip() == "12.3 Evidências técnicas":
        paragrafo.paragraph_format.page_break_before = True
        break
doc.save(passo)


executivo = RAIZ / "RelatorioExecutivoConsolidado_GeoAI_Mentor.docx"
doc = Document(executivo)
itens = {
    "Instale as dependências com pip install -r requirements.txt.",
    "Execute os testes com python -m pytest -q.",
    "Inicie a interface com streamlit run streamlit_app.py.",
}
for paragrafo in doc.paragraphs:
    if paragrafo.text.strip() in itens:
        paragrafo.style = "List Bullet"
doc.save(executivo)
